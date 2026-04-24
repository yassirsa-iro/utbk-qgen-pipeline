"""
01_extract.py — Step 1: Ekstraksi Teks dari Input Soal
=======================================================
Mendukung format: .docx, .pdf, .png, .jpg
Rumus matematika → dikonversi ke LaTeX via Claude Vision API
Output: /data/raw_extracted/<nama_file>.json
"""

import os
import json
import base64
import anthropic
from pathlib import Path
from datetime import datetime

import docx
import pdfplumber
from PIL import Image

# ── Konfigurasi Path ──────────────────────────────────────────────────────────
INPUT_DIR  = Path("input_raw")
OUTPUT_DIR = Path("data/raw_extracted")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Client Anthropic (Claude) ─────────────────────────────────────────────────
client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

PROMPT_VISION = """
Kamu adalah asisten transkripsi soal ujian UTBK/SNBT (TPS dan TKA).

Tugasmu: Transkripsi SELURUH isi gambar ini menjadi teks terstruktur.

Aturan penting:
1. Setiap soal diawali dengan nomor soal, contoh: **Soal 1**
2. Semua rumus matematika, fisika, atau kimia WAJIB ditulis dalam format LaTeX.
   - Contoh inline: $x^2 + y^2 = r^2$
   - Contoh display: $$\frac{d}{dx}[x^n] = nx^{n-1}$$
3. Pilihan jawaban ditulis sebagai: A) ... B) ... C) ... D) ... E) ...
4. Jika ada tabel, transkripsi sebagai Markdown table.
5. Jika ada grafik/gambar yang tidak bisa ditranskripsi, tulis: [GAMBAR: deskripsi singkat]
6. Jangan tambahkan komentar atau penjelasan, langsung transkripsi saja.
"""

# ── Fungsi Ekstraksi ──────────────────────────────────────────────────────────

def extract_docx(path: Path) -> str:
    """Ekstrak teks dari file Word (.docx)"""
    doc = docx.Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Ekstrak teks dari tabel juga
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def extract_pdf(path: Path) -> str:
    """Ekstrak teks dari PDF digital (bukan scan)"""
    pages_text = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(f"--- Halaman {i} ---\n{text.strip()}")
    return "\n\n".join(pages_text)


def extract_image_via_claude(path: Path) -> str:
    """Ekstrak teks dari gambar (PNG/JPG) menggunakan Claude Vision"""
    with open(path, "rb") as f:
        image_data = base64.standard_b64encode(f.read()).decode("utf-8")

    suffix = path.suffix.lower()
    media_type_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg"}
    media_type = media_type_map.get(suffix, "image/png")

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=4096,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_data,
                        },
                    },
                    {"type": "text", "text": PROMPT_VISION},
                ],
            }
        ],
    )
    return message.content[0].text


def pdf_is_scanned(path: Path) -> bool:
    """Cek apakah PDF adalah hasil scan (tidak ada teks yang bisa diekstrak)"""
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages[:3]:  # Cek 3 halaman pertama
            text = page.extract_text()
            if text and len(text.strip()) > 50:
                return False
    return True


def extract_pdf_scan(path: Path) -> str:
    """Untuk PDF scan: konversi setiap halaman ke gambar, lalu OCR via Claude"""
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    all_text = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=200)
        img_path = OUTPUT_DIR / f"_temp_page_{page_num}.png"
        pix.save(str(img_path))
        print(f"  → OCR halaman {page_num + 1}/{len(doc)}...")
        text = extract_image_via_claude(img_path)
        all_text.append(f"--- Halaman {page_num + 1} ---\n{text}")
        img_path.unlink()  # Hapus file temp

    doc.close()
    return "\n\n".join(all_text)


# ── Main ──────────────────────────────────────────────────────────────────────

def process_file(file_path: Path) -> dict:
    """Proses satu file dan return hasilnya sebagai dict"""
    suffix = file_path.suffix.lower()
    print(f"\n📄 Memproses: {file_path.name}")

    if suffix == ".docx":
        print("  → Format: Word (.docx)")
        raw_text = extract_docx(file_path)
        method = "python-docx"

    elif suffix == ".pdf":
        if pdf_is_scanned(file_path):
            print("  → Format: PDF Scan (OCR via Claude Vision)")
            raw_text = extract_pdf_scan(file_path)
            method = "claude-vision-ocr"
        else:
            print("  → Format: PDF Digital")
            raw_text = extract_pdf(file_path)
            method = "pdfplumber"

    elif suffix in [".png", ".jpg", ".jpeg"]:
        print("  → Format: Gambar (OCR via Claude Vision)")
        raw_text = extract_image_via_claude(file_path)
        method = "claude-vision-ocr"

    else:
        print(f"  ⚠️  Format tidak didukung: {suffix}, dilewati.")
        return None

    result = {
        "file_asal": file_path.name,
        "format": suffix,
        "metode_ekstraksi": method,
        "waktu_ekstraksi": datetime.now().isoformat(),
        "teks_mentah": raw_text,
        "status": "extracted"
    }

    # Simpan ke output
    output_filename = OUTPUT_DIR / f"{file_path.stem}.json"
    with open(output_filename, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"  ✅ Tersimpan: {output_filename}")
    return result


def main():
    print("=" * 60)
    print("  Step 1: Ekstraksi Teks dari Input Soal")
    print("=" * 60)

    supported_ext = {".docx", ".pdf", ".png", ".jpg", ".jpeg"}
    files = [
        f for f in INPUT_DIR.iterdir()
        if f.is_file() and f.suffix.lower() in supported_ext
    ]

    if not files:
        print("\n⚠️  Tidak ada file baru di input_raw/. Pipeline selesai untuk step ini.")
        return

    print(f"\n📂 Ditemukan {len(files)} file untuk diproses.\n")
    processed = 0
    for file_path in files:
        result = process_file(file_path)
        if result:
            processed += 1

    print(f"\n{'='*60}")
    print(f"  ✅ Step 1 Selesai — {processed}/{len(files)} file berhasil diekstrak.")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
